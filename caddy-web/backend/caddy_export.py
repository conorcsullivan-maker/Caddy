"""
Conversation export — render a Caddy chat (archived or active) as a Word
document. Used by the owner-download endpoints so Conor and Drew can share
real session transcripts back and forth during beta.
"""
import io
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


FOREST = RGBColor(0x1F, 0x4E, 0x44)
MUTED = RGBColor(0x77, 0x77, 0x77)
INK = RGBColor(0x1A, 0x1A, 0x1A)


def _fmt_dt(iso: Optional[str]) -> str:
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        return dt.strftime("%b %d, %Y · %-I:%M %p")
    except Exception:
        return iso


def _add_meta_line(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.font.size = Pt(10)
    run_label.font.bold = True
    run_label.font.color.rgb = MUTED
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = INK


def conversation_to_docx_bytes(
    *,
    full_name: str,
    username: str,
    kind: str,
    course_name: Optional[str],
    total_score: Optional[int],
    started_at: Optional[str],
    ended_at: Optional[str],
    round_metadata: Optional[dict],
    messages: list[dict],
    is_active: bool = False,
) -> bytes:
    """Render a single conversation into a .docx blob.

    `messages` is the chat history list-of-dicts with role/content shape that
    we already store in the DB. `is_active` flips the document framing so it
    reads as an in-progress chat rather than an archived round/casual."""
    doc = Document()

    # 1-inch margins all around
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title.add_run(f"Caddy conversation — {full_name}")
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = FOREST

    # Subtitle / kind label
    sub = doc.add_paragraph()
    if is_active:
        kind_label = "Active conversation (in progress)"
    elif kind == "round":
        kind_label = "Logged round"
    else:
        kind_label = "Casual chat"
    s_run = sub.add_run(kind_label)
    s_run.font.size = Pt(11)
    s_run.font.italic = True
    s_run.font.color.rgb = MUTED

    # Metadata block
    if course_name:
        _add_meta_line(doc, "Course", course_name)
    if total_score is not None:
        _add_meta_line(doc, "Score", str(total_score))
    if round_metadata:
        diff = round_metadata.get("differential")
        if diff is not None:
            _add_meta_line(doc, "Differential", f"{diff:.1f}")
        hcp = round_metadata.get("handicap_after")
        if hcp is not None:
            _add_meta_line(doc, "Handicap after", str(hcp))
    if started_at:
        _add_meta_line(doc, "Started", _fmt_dt(started_at))
    if ended_at and not is_active:
        _add_meta_line(doc, "Ended", _fmt_dt(ended_at))
    _add_meta_line(doc, "Player", f"{full_name} (@{username})")

    # Divider
    doc.add_paragraph()
    rule = doc.add_paragraph()
    r_run = rule.add_run("─" * 60)
    r_run.font.color.rgb = MUTED
    r_run.font.size = Pt(9)

    # Hole-by-hole scorecard for round conversations
    if round_metadata and isinstance(round_metadata.get("hole_scores"), list):
        hole_scores = round_metadata["hole_scores"]
        if any(s is not None for s in hole_scores):
            head = doc.add_paragraph()
            h_run = head.add_run("Scorecard")
            h_run.font.bold = True
            h_run.font.size = Pt(12)
            h_run.font.color.rgb = FOREST

            line = doc.add_paragraph()
            cells = []
            for i, s in enumerate(hole_scores[:18]):
                cells.append(f"H{i+1}: {s if s is not None else '—'}")
            text_run = line.add_run("   ".join(cells))
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = INK
            doc.add_paragraph()

    # Messages
    if not messages:
        empty = doc.add_paragraph()
        empty_run = empty.add_run("(No messages in this conversation.)")
        empty_run.font.italic = True
        empty_run.font.color.rgb = MUTED
    else:
        for m in messages:
            role = m.get("role") or "unknown"
            content = (m.get("content") or "").strip()
            if not content:
                continue
            speaker = "You" if role == "user" else "Caddy"
            color = INK if role == "user" else FOREST

            speaker_p = doc.add_paragraph()
            sp_run = speaker_p.add_run(speaker)
            sp_run.font.bold = True
            sp_run.font.size = Pt(11)
            sp_run.font.color.rgb = color

            body_p = doc.add_paragraph()
            body_p.paragraph_format.left_indent = Inches(0.25)
            body_p.paragraph_format.space_after = Pt(8)
            for j, line in enumerate(content.split("\n")):
                if j > 0:
                    body_p.add_run("\n")
                br = body_p.add_run(line)
                br.font.size = Pt(11)
                br.font.color.rgb = INK

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = foot.add_run(
        f"Exported {datetime.utcnow().strftime('%b %d, %Y')} · Caddy beta"
    )
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(prefix: str, course_name: Optional[str], date_iso: Optional[str]) -> str:
    """Build a download filename like 'Caddy_Round_ButterBrook_2026-05-21.docx'."""
    parts = [prefix]
    if course_name:
        cleaned = "".join(c for c in course_name if c.isalnum() or c in (" ", "-", "_")).strip()
        cleaned = cleaned.replace(" ", "")
        if cleaned:
            parts.append(cleaned[:40])
    if date_iso:
        parts.append(date_iso[:10])
    return "_".join(parts) + ".docx"
